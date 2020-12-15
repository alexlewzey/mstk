"""
Provides a selection of tools which take python data structures and objects and add them to a word document using the
python-docx module
"""

import logging
from typing import List, Optional, Callable
import pandas as pd
from docx import Document, shared
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def show_styles(doc: Document) -> None:
    """show all doc styles"""
    for s in doc.styles:
        print(s.name)


class DocStyles:
    """document style names that can be assigned to different document paragraphs and components"""
    # text styles
    heading5 = 'Heading 5'
    heading4 = 'Heading 4'
    heading3 = 'Heading 3'
    heading2 = 'Heading 2'
    heading1 = 'Heading 1'
    normal = 'Normal'
    list_bullet = 'List Bullet'
    header = 'Header'
    footer = 'Footer'

    # table styles
    table = 'Normal Table'  # 'Grid Table 5 Dark - Accent 3'

    # # table text styles
    # tableTitle = 'Table Title'
    # tableSource = 'Table Source'
    # tableText = 'Table Text'
    # tableNumber = 'Table Number'
    # tableHeading = 'Table Heading'
    #
    # # figure text styles
    # figureTitle = 'Figure Title'
    # normalFigures = 'NormalFigures'
    # figureSource = 'Figure Source'


class FootNotes:
    """footnotes and sources in figures and tables"""
    NOTE = 'something'


def add_table_from_list(
    doc: Document, title: str, columns: List[str], rows: List[str], data: List[List],
    footnote: Optional[str] = None) -> Document:
    """
    add a table with row and column headers to a word document and return the document
    Args:
        doc: python-docx object
        title: table title
        columns: column headers
        rows: row headers
        data: list of lists [row, row, row]
        footnote: footnote string

    Returns:
        python-docx object with table added to the end of the doc
    """
    logger.info('Adding table from list to doc')
    # adding title
    doc.add_paragraph(text=title, style=DocStyles.footer)

    # adding table
    no_rows = len(data)
    no_cols = len(columns)
    assert no_cols == len(data[0]), 'The number of headers and columns of data are not equal'
    assert no_rows == len(data), 'The number of row headers and rows of data are not equal'

    table = doc.add_table(no_rows + 1, no_cols + 1, DocStyles.table)

    # populating column headers
    for j in range(no_cols):
        table.cell(0, j + 1).text = columns[j]
        table.cell(0, j + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # populating row headers
    for i in range(no_rows):
        table.cell(i + 1, 0).text = rows[i]
        table.cell(i + 1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # populating data cells
    for i in range(no_rows):
        for j in range(no_cols):
            table.cell(i + 1, j + 1).text = str(data[i][j])
            table.cell(i + 1, j + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if footnote:
        # adding a footnote
        doc.add_paragraph(text=footnote, style=DocStyles.normal)
    logger.info('Finished adding table from list to doc')

    return doc


def add_table_from_df(doc: Document, title: str, df: pd.DataFrame, footnote: Optional[str] = None) -> Document:
    """
    add a table to the doc from a DataFrame and return the doc
    Args:
        doc: document object
        title: table title
        df: DataFrame
        footnote: table footnote

    Returns:
        populated document
    """
    rows = df.index.tolist()
    columns = df.columns.tolist()
    values = [list(row) for row in df.values]
    return add_table_from_list(doc=doc,
                               title=title,
                               columns=columns,
                               rows=rows,
                               data=values,
                               footnote=footnote)


def add_figure(doc: Document, title: str, img_path: str, source: Optional[str] = None) -> Document:
    """
    add a figure to a word document and return the document
    Args:
        doc: document object
        title: table title
        img_path: path to image file
        source: footnote source of image

    Returns:
        populated document
    """
    # adding a figure table
    fig_table = doc.add_table(2, 1)

    # figure title
    fig_table.cell(0, 0).text = title
    fig_table.cell(0, 0).paragraphs[0].style = DocStyles.normal

    # figure image
    para = fig_table.cell(1, 0).paragraphs[0]
    para.style = DocStyles.normal
    run = para.add_run()
    run.add_picture(img_path, width=shared.Cm(15.2))

    # figure source
    # fig_table.cell(2, 0).text = source
    # fig_table.cell(2, 0).paragraphs[0].style = DocStyles.normal
    if source:
        # adding a footnote
        doc.add_paragraph(text=source, style=DocStyles.normal)
    logger.info(f'Added figure: {title}')
    return doc


# tools for formatting text

def num_round(x) -> str:
    """comma separated with only the largest 2 units as non-zero"""
    if x > 100:
        x = round(x, 2 - len(str(int(x))))
    elif x < 100:
        x = round(x, 3 - len(str(int(x))))
    elif abs(x) > 10:
        x = round(x, -1)
    else:
        x = x
    x = int(round(x))
    return f'{x:,}'


def fmt_written_list_with_values_in_bracket(names: List[str], values: List[str], length: int = 3) -> str:
    """take in two lists and convert into a string list eg Derby (40%), Matlock (30%) and Winster (20%)"""
    combined = []
    for name, val in zip(names, values):
        combined.append(f'{name} ({val})')
    return ', '.join(combined[:length]) + '.'


def fmt_written_comma_sep_list(lst: List, preprocessing: Optional[Callable] = None) -> str:
    """format a list as a comma separated string with an and in the final gap"""
    if preprocessing:
        lst = [preprocessing(x) for x in lst]
    return ', '.join([str(x) for x in lst[:-1]]) + ' and ' + str(lst[-1])
